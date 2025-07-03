#!/usr/bin/env python3
"""
Test script for the Document Comparison Tool

This script creates sample documents and tests the comparison functionality.
"""

import tempfile
import os
from pathlib import Path
import docx
from document_compare import DocumentComparisonTool

def create_sample_documents():
    """Create sample Word documents for testing."""
    
    # Create temporary directory
    temp_dir = Path(tempfile.mkdtemp())
    
    # Create baseline template document
    baseline_doc = docx.Document()
    baseline_doc.add_heading('Project Requirements Template', 0)
    baseline_doc.add_paragraph('This is a template document for project requirements.')
    baseline_doc.add_heading('1. Introduction', level=1)
    baseline_doc.add_paragraph('The project must include the following requirements:')
    baseline_doc.add_paragraph('• Requirement 1: System must be secure')
    baseline_doc.add_paragraph('• Requirement 2: System must be scalable')
    baseline_doc.add_paragraph('• Requirement 3: System must be reliable')
    baseline_doc.add_heading('2. Technical Specifications', level=1)
    baseline_doc.add_paragraph('Technical details shall be documented here.')
    baseline_doc.add_heading('3. Timeline', level=1)
    baseline_doc.add_paragraph('Project timeline and milestones.')
    
    baseline_path = temp_dir / "template.docx"
    baseline_doc.save(baseline_path)
    
    # Create comparison document (similar but with differences)
    comparison_doc = docx.Document()
    comparison_doc.add_heading('Project Requirements Document', 0)  # Different title
    comparison_doc.add_paragraph('This is a document for project requirements.')  # Missing "template"
    comparison_doc.add_heading('1. Introduction', level=1)
    comparison_doc.add_paragraph('The project must include the following requirements:')
    comparison_doc.add_paragraph('• Requirement 1: System must be secure')
    comparison_doc.add_paragraph('• Requirement 2: System must be scalable')
    # Missing Requirement 3
    comparison_doc.add_paragraph('• Requirement 4: System must be user-friendly')  # Added requirement
    comparison_doc.add_heading('2. Technical Specifications', level=1)
    comparison_doc.add_paragraph('Technical details shall be documented here.')
    comparison_doc.add_heading('3. Timeline', level=1)
    comparison_doc.add_paragraph('Project timeline and milestones.')
    comparison_doc.add_heading('4. Budget', level=1)  # Added section
    comparison_doc.add_paragraph('Budget considerations and cost estimates.')
    
    comparison_path = temp_dir / "comparison.docx"
    comparison_doc.save(comparison_path)
    
    return baseline_path, comparison_path, temp_dir

def test_basic_functionality():
    """Test basic comparison functionality."""
    print("Testing basic document comparison functionality...")
    
    try:
        # Create sample documents
        baseline_path, comparison_path, temp_dir = create_sample_documents()
        
        # Create comparison tool
        tool = DocumentComparisonTool(log_level='WARNING')  # Quiet mode for testing
        
        # Perform comparison
        result = tool.compare_documents(
            baseline_path,
            comparison_path,
            temp_dir / "test_reports"
        )
        
        # Validate results
        assert result.similarity_ratio > 0.0, "Similarity ratio should be greater than 0"
        assert result.similarity_ratio < 1.0, "Similarity ratio should be less than 1 (documents are different)"
        assert len(result.differences) > 0, "Should detect differences"
        assert len(result.missing_content) > 0, "Should detect missing content"
        assert len(result.added_content) > 0, "Should detect added content"
        
        # Check that reports were generated
        report_dir = temp_dir / "test_reports"
        html_files = list(report_dir.glob("*.html"))
        json_files = list(report_dir.glob("*.json"))
        csv_files = list(report_dir.glob("*.csv"))
        
        assert len(html_files) > 0, "HTML report should be generated"
        assert len(json_files) > 0, "JSON report should be generated"
        assert len(csv_files) > 0, "CSV report should be generated"
        
        print("✅ Basic functionality test passed!")
        print(f"   Similarity: {result.similarity_ratio:.2%}")
        print(f"   Differences: {len(result.differences)}")
        print(f"   Missing content: {len(result.missing_content)}")
        print(f"   Added content: {len(result.added_content)}")
        print(f"   Reports saved to: {report_dir}")
        
        # Cleanup
        import shutil
        shutil.rmtree(temp_dir)
        
        return True
        
    except Exception as e:
        print(f"❌ Test failed: {str(e)}")
        return False

def test_error_handling():
    """Test error handling for various edge cases."""
    print("\nTesting error handling...")
    
    try:
        tool = DocumentComparisonTool(log_level='WARNING')
        
        # Test with non-existent files
        try:
            result = tool.compare_documents("nonexistent1.docx", "nonexistent2.docx")
            print("❌ Should have raised an error for non-existent files")
            return False
        except Exception:
            print("✅ Correctly handled non-existent files")
        
        # Test with unsupported file type
        temp_dir = Path(tempfile.mkdtemp())
        dummy_file = temp_dir / "test.txt"
        dummy_file.write_text("This is a text file")
        
        try:
            result = tool.compare_documents(dummy_file, dummy_file)
            print("❌ Should have raised an error for unsupported file type")
            return False
        except Exception:
            print("✅ Correctly handled unsupported file type")
        
        # Cleanup
        import shutil
        shutil.rmtree(temp_dir)
        
        return True
        
    except Exception as e:
        print(f"❌ Error handling test failed: {str(e)}")
        return False

def test_cli_interface():
    """Test the command-line interface."""
    print("\nTesting CLI interface...")
    
    try:
        # Create sample documents
        baseline_path, comparison_path, temp_dir = create_sample_documents()
        
        # Test CLI help
        import subprocess
        import sys
        
        result = subprocess.run([
            sys.executable, "document_compare.py", "--help"
        ], capture_output=True, text=True)
        
        if result.returncode == 0 and "usage:" in result.stdout.lower():
            print("✅ CLI help works correctly")
        else:
            print("❌ CLI help test failed")
            return False
        
        # Test basic CLI usage
        output_dir = temp_dir / "cli_test_reports"
        result = subprocess.run([
            sys.executable, "document_compare.py",
            str(baseline_path),
            str(comparison_path),
            "--output-dir", str(output_dir),
            "--log-level", "WARNING"
        ], capture_output=True, text=True)
        
        if result.returncode == 0:
            print("✅ CLI comparison works correctly")
            
            # Check if reports were generated
            if output_dir.exists() and list(output_dir.glob("*.html")):
                print("✅ CLI generated reports correctly")
            else:
                print("❌ CLI did not generate expected reports")
                return False
        else:
            print(f"❌ CLI comparison failed: {result.stderr}")
            return False
        
        # Cleanup
        import shutil
        shutil.rmtree(temp_dir)
        
        return True
        
    except Exception as e:
        print(f"❌ CLI test failed: {str(e)}")
        return False

def main():
    """Run all tests."""
    print("Document Comparison Tool - Test Suite")
    print("=" * 50)
    
    tests = [
        test_basic_functionality,
        test_error_handling,
        test_cli_interface
    ]
    
    passed = 0
    total = len(tests)
    
    for test_func in tests:
        if test_func():
            passed += 1
    
    print("\n" + "=" * 50)
    print(f"Test Results: {passed}/{total} tests passed")
    
    if passed == total:
        print("🎉 All tests passed! The document comparison tool is working correctly.")
    else:
        print("⚠️  Some tests failed. Please check the errors above.")
    
    return passed == total

if __name__ == "__main__":
    success = main()
    exit(0 if success else 1) 