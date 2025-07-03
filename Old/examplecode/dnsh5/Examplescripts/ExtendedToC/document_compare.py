#!/usr/bin/env python3
"""
Document Comparison Tool

A comprehensive tool for comparing documents against a baseline template.
Supports Word (.docx) and PDF formats with detailed difference reporting.

Author: Your Name
Date: 2025
"""

import argparse
import logging
import sys
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union
import json
from dataclasses import dataclass, asdict
from datetime import datetime

# Third-party imports will be available after installing requirements
try:
    import docx
    import pdfplumber
    import difflib
    from jinja2 import Template
    import pandas as pd
except ImportError as e:
    print(f"Missing required package: {e}")
    print("Please install requirements: pip install -r requirements_document_compare.txt")
    sys.exit(1)


@dataclass
class DocumentContent:
    """Data class for document content and metadata."""
    file_path: str
    file_type: str
    text_content: str
    paragraphs: List[str]
    word_count: int
    extraction_success: bool
    error_message: Optional[str] = None


@dataclass
class ComparisonResult:
    """Data class for comparison results."""
    baseline_file: str
    comparison_file: str
    similarity_ratio: float
    differences: List[Dict]
    missing_content: List[str]
    added_content: List[str]
    timestamp: str


class DocumentParser:
    """Handles parsing of different document formats."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def parse_document(self, file_path: Path) -> DocumentContent:
        """
        Parse a document and extract text content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            DocumentContent object with extracted text and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return DocumentContent(
                file_path=str(file_path),
                file_type="unknown",
                text_content="",
                paragraphs=[],
                word_count=0,
                extraction_success=False,
                error_message=f"File not found: {file_path}"
            )
        
        try:
            if file_path.suffix.lower() == '.docx':
                return self._parse_docx(file_path)
            elif file_path.suffix.lower() == '.pdf':
                return self._parse_pdf(file_path)
            else:
                return DocumentContent(
                    file_path=str(file_path),
                    file_type=file_path.suffix.lower(),
                    text_content="",
                    paragraphs=[],
                    word_count=0,
                    extraction_success=False,
                    error_message=f"Unsupported file type: {file_path.suffix}"
                )
        except Exception as e:
            self.logger.error(f"Error parsing {file_path}: {str(e)}")
            return DocumentContent(
                file_path=str(file_path),
                file_type=file_path.suffix.lower(),
                text_content="",
                paragraphs=[],
                word_count=0,
                extraction_success=False,
                error_message=str(e)
            )
    
    def _parse_docx(self, file_path: Path) -> DocumentContent:
        """Parse Word document."""
        doc = docx.Document(file_path)
        paragraphs = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # Only add non-empty paragraphs
                paragraphs.append(text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(f"[TABLE] {cell_text}")
        
        full_text = '\n'.join(paragraphs)
        word_count = len(full_text.split())
        
        return DocumentContent(
            file_path=str(file_path),
            file_type=".docx",
            text_content=full_text,
            paragraphs=paragraphs,
            word_count=word_count,
            extraction_success=True
        )
    
    def _parse_pdf(self, file_path: Path) -> DocumentContent:
        """Parse PDF document."""
        paragraphs = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text:
                    # Split into paragraphs and clean up
                    page_paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
                    paragraphs.extend(page_paragraphs)
        
        full_text = '\n'.join(paragraphs)
        word_count = len(full_text.split())
        
        return DocumentContent(
            file_path=str(file_path),
            file_type=".pdf",
            text_content=full_text,
            paragraphs=paragraphs,
            word_count=word_count,
            extraction_success=True
        )


class DocumentComparator:
    """Handles document comparison logic."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def compare_documents(self, baseline: DocumentContent, comparison: DocumentContent) -> ComparisonResult:
        """
        Compare two documents and generate detailed differences.
        
        Args:
            baseline: The baseline/template document
            comparison: The document to compare against baseline
            
        Returns:
            ComparisonResult with detailed comparison data
        """
        if not baseline.extraction_success:
            raise ValueError(f"Baseline document parsing failed: {baseline.error_message}")
        
        if not comparison.extraction_success:
            raise ValueError(f"Comparison document parsing failed: {comparison.error_message}")
        
        # Calculate overall similarity
        similarity_ratio = difflib.SequenceMatcher(
            None, baseline.text_content, comparison.text_content
        ).ratio()
        
        # Get detailed differences
        differences = self._get_detailed_differences(baseline, comparison)
        
        # Find missing and added content
        missing_content = self._find_missing_content(baseline, comparison)
        added_content = self._find_added_content(baseline, comparison)
        
        return ComparisonResult(
            baseline_file=baseline.file_path,
            comparison_file=comparison.file_path,
            similarity_ratio=similarity_ratio,
            differences=differences,
            missing_content=missing_content,
            added_content=added_content,
            timestamp=datetime.now().isoformat()
        )
    
    def _get_detailed_differences(self, baseline: DocumentContent, comparison: DocumentContent) -> List[Dict]:
        """Get detailed line-by-line differences."""
        baseline_lines = baseline.text_content.split('\n')
        comparison_lines = comparison.text_content.split('\n')
        
        differ = difflib.unified_diff(
            baseline_lines, 
            comparison_lines,
            fromfile='baseline',
            tofile='comparison',
            lineterm=''
        )
        
        differences = []
        for line in differ:
            if line.startswith('@@'):
                continue
            elif line.startswith('-'):
                differences.append({
                    'type': 'removed',
                    'content': line[1:],
                    'line_type': 'baseline'
                })
            elif line.startswith('+'):
                differences.append({
                    'type': 'added',
                    'content': line[1:],
                    'line_type': 'comparison'
                })
        
        return differences
    
    def _find_missing_content(self, baseline: DocumentContent, comparison: DocumentContent) -> List[str]:
        """Find content present in baseline but missing in comparison."""
        baseline_paragraphs = set(baseline.paragraphs)
        comparison_paragraphs = set(comparison.paragraphs)
        return list(baseline_paragraphs - comparison_paragraphs)
    
    def _find_added_content(self, baseline: DocumentContent, comparison: DocumentContent) -> List[str]:
        """Find content present in comparison but not in baseline."""
        baseline_paragraphs = set(baseline.paragraphs)
        comparison_paragraphs = set(comparison.paragraphs)
        return list(comparison_paragraphs - baseline_paragraphs)


class ReportGenerator:
    """Generates comparison reports in various formats."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def generate_html_report(self, result: ComparisonResult, output_path: Path) -> None:
        """Generate an HTML report with highlighted differences."""
        html_template = """
<!DOCTYPE html>
<html>
<head>
    <title>Document Comparison Report</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 20px; }
        .header { background-color: #f0f0f0; padding: 15px; border-radius: 5px; }
        .summary { margin: 20px 0; }
        .stat-box { display: inline-block; margin: 10px; padding: 10px; background-color: #e8f4fd; border-radius: 5px; }
        .differences { margin: 20px 0; }
        .diff-removed { background-color: #ffe6e6; padding: 5px; margin: 2px 0; border-left: 4px solid #ff0000; }
        .diff-added { background-color: #e6ffe6; padding: 5px; margin: 2px 0; border-left: 4px solid #00aa00; }
        .content-section { margin: 20px 0; padding: 15px; border: 1px solid #ddd; border-radius: 5px; }
        .content-item { margin: 5px 0; padding: 8px; background-color: #f9f9f9; border-radius: 3px; }
        pre { white-space: pre-wrap; word-wrap: break-word; }
    </style>
</head>
<body>
    <div class="header">
        <h1>Document Comparison Report</h1>
        <p><strong>Generated:</strong> {{ result.timestamp }}</p>
        <p><strong>Baseline:</strong> {{ result.baseline_file }}</p>
        <p><strong>Comparison:</strong> {{ result.comparison_file }}</p>
    </div>
    
    <div class="summary">
        <h2>Summary</h2>
        <div class="stat-box">
            <strong>Similarity:</strong> {{ "%.2f"|format(result.similarity_ratio * 100) }}%
        </div>
        <div class="stat-box">
            <strong>Differences:</strong> {{ result.differences|length }}
        </div>
        <div class="stat-box">
            <strong>Missing Content:</strong> {{ result.missing_content|length }} items
        </div>
        <div class="stat-box">
            <strong>Added Content:</strong> {{ result.added_content|length }} items
        </div>
    </div>
    
    {% if result.differences %}
    <div class="differences">
        <h2>Line-by-Line Differences</h2>
        {% for diff in result.differences %}
            {% if diff.type == 'removed' %}
                <div class="diff-removed">
                    <strong>REMOVED:</strong> {{ diff.content }}
                </div>
            {% elif diff.type == 'added' %}
                <div class="diff-added">
                    <strong>ADDED:</strong> {{ diff.content }}
                </div>
            {% endif %}
        {% endfor %}
    </div>
    {% endif %}
    
    {% if result.missing_content %}
    <div class="content-section">
        <h2>Missing Content (Present in baseline, absent in comparison)</h2>
        {% for item in result.missing_content %}
            <div class="content-item">{{ item }}</div>
        {% endfor %}
    </div>
    {% endif %}
    
    {% if result.added_content %}
    <div class="content-section">
        <h2>Added Content (Present in comparison, absent in baseline)</h2>
        {% for item in result.added_content %}
            <div class="content-item">{{ item }}</div>
        {% endfor %}
    </div>
    {% endif %}
</body>
</html>
        """
        
        template = Template(html_template)
        html_content = template.render(result=result)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        self.logger.info(f"HTML report generated: {output_path}")
    
    def generate_json_report(self, result: ComparisonResult, output_path: Path) -> None:
        """Generate a structured JSON report."""
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(asdict(result), f, indent=2, ensure_ascii=False)
        
        self.logger.info(f"JSON report generated: {output_path}")
    
    def generate_csv_report(self, result: ComparisonResult, output_path: Path) -> None:
        """Generate a CSV report for detailed analysis."""
        # Create a DataFrame with all differences
        data = []
        
        # Add basic info
        data.append({
            'Type': 'Summary',
            'Category': 'Similarity',
            'Content': f"{result.similarity_ratio:.2%}",
            'Source': 'Calculated'
        })
        
        # Add differences
        for diff in result.differences:
            data.append({
                'Type': 'Difference',
                'Category': diff['type'],
                'Content': diff['content'],
                'Source': diff['line_type']
            })
        
        # Add missing content
        for content in result.missing_content:
            data.append({
                'Type': 'Missing',
                'Category': 'baseline_only',
                'Content': content,
                'Source': 'baseline'
            })
        
        # Add added content
        for content in result.added_content:
            data.append({
                'Type': 'Added',
                'Category': 'comparison_only',
                'Content': content,
                'Source': 'comparison'
            })
        
        df = pd.DataFrame(data)
        df.to_csv(output_path, index=False, encoding='utf-8')
        
        self.logger.info(f"CSV report generated: {output_path}")


class DocumentComparisonTool:
    """Main application class orchestrating the document comparison process."""
    
    def __init__(self, log_level: str = 'INFO'):
        self.setup_logging(log_level)
        self.logger = logging.getLogger(__name__)
        
        self.parser = DocumentParser()
        self.comparator = DocumentComparator()
        self.reporter = ReportGenerator()
    
    def setup_logging(self, log_level: str) -> None:
        """Setup logging configuration."""
        logging.basicConfig(
            level=getattr(logging, log_level.upper()),
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            handlers=[
                logging.StreamHandler(sys.stdout),
                logging.FileHandler('document_compare.log')
            ]
        )
    
    def compare_documents(
        self, 
        baseline_path: Union[str, Path], 
        comparison_path: Union[str, Path],
        output_dir: Union[str, Path] = None
    ) -> ComparisonResult:
        """
        Main method to compare two documents.
        
        Args:
            baseline_path: Path to the baseline/template document
            comparison_path: Path to the document to compare
            output_dir: Directory to save reports (optional)
            
        Returns:
            ComparisonResult object
        """
        baseline_path = Path(baseline_path)
        comparison_path = Path(comparison_path)
        
        self.logger.info(f"Starting comparison: {baseline_path} vs {comparison_path}")
        
        # Parse documents
        baseline_content = self.parser.parse_document(baseline_path)
        comparison_content = self.parser.parse_document(comparison_path)
        
        # Perform comparison
        result = self.comparator.compare_documents(baseline_content, comparison_content)
        
        # Generate reports if output directory is specified
        if output_dir:
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_name = f"comparison_report_{timestamp}"
            
            # Generate all report formats
            self.reporter.generate_html_report(
                result, 
                output_dir / f"{base_name}.html"
            )
            self.reporter.generate_json_report(
                result, 
                output_dir / f"{base_name}.json"
            )
            self.reporter.generate_csv_report(
                result, 
                output_dir / f"{base_name}.csv"
            )
        
        return result


def main():
    """Command-line interface for the document comparison tool."""
    parser = argparse.ArgumentParser(
        description="Compare documents against a baseline template",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python document_compare.py template.docx document.docx
  python document_compare.py template.docx document.pdf --output-dir reports/
  python document_compare.py template.docx document.docx --log-level DEBUG
        """
    )
    
    parser.add_argument(
        'baseline',
        help='Path to the baseline/template document'
    )
    parser.add_argument(
        'comparison',
        help='Path to the document to compare'
    )
    parser.add_argument(
        '--output-dir', '-o',
        help='Directory to save comparison reports'
    )
    parser.add_argument(
        '--log-level', '-l',
        choices=['DEBUG', 'INFO', 'WARNING', 'ERROR'],
        default='INFO',
        help='Set the logging level'
    )
    
    args = parser.parse_args()
    
    try:
        # Create comparison tool
        tool = DocumentComparisonTool(log_level=args.log_level)
        
        # Perform comparison
        result = tool.compare_documents(
            args.baseline,
            args.comparison,
            args.output_dir
        )
        
        # Print summary to console
        print("\n" + "="*50)
        print("DOCUMENT COMPARISON SUMMARY")
        print("="*50)
        print(f"Baseline: {result.baseline_file}")
        print(f"Comparison: {result.comparison_file}")
        print(f"Similarity: {result.similarity_ratio:.2%}")
        print(f"Differences found: {len(result.differences)}")
        print(f"Missing content items: {len(result.missing_content)}")
        print(f"Added content items: {len(result.added_content)}")
        
        if args.output_dir:
            print(f"\nReports saved to: {args.output_dir}")
        
        print("="*50)
        
    except Exception as e:
        logging.error(f"Comparison failed: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main() 